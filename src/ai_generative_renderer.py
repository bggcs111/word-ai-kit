"""
AI生成式文档渲染器
渲染AI生成的新文档结构，统一格式（字体、间距）
"""
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import tempfile
import os


class AIGenerativeRenderer:
    """AI生成式文档渲染器"""
    
    def __init__(self):
        """初始化渲染器"""
        self.default_font = "宋体"
        self.default_font_size = Pt(12)
        self.heading_font = "黑体"
        self.heading_sizes = {
            1: Pt(16),
            2: Pt(14),
            3: Pt(12)
        }
        self.line_spacing = 1.5
    
    def render(
        self,
        generated_content: Dict[str, Any],
        output_path: str
    ) -> str:
        """
        渲染AI生成的文档
        
        Args:
            generated_content: AI生成的内容结构
            output_path: 输出文件路径
            
        Returns:
            输出文件路径
        """
        structure = generated_content.get('structure', [])
        tables_data = generated_content.get('tables', [])
        images_data = generated_content.get('images', [])
        
        if not structure:
            raise ValueError("生成的内容结构为空")
        
        # 创建新文档
        doc = Document()
        
        # 设置默认样式
        self._setup_default_styles(doc)
        
        # 渲染结构
        for item in structure:
            item_type = item.get('type')
            
            if item_type == 'heading':
                self._render_heading(doc, item)
            elif item_type == 'paragraph':
                self._render_paragraph(doc, item)
            elif item_type == 'table':
                self._render_table(doc, item, tables_data)
            elif item_type == 'image':
                self._render_image(doc, item, images_data)
        
        # 保存文档
        doc.save(output_path)
        
        return output_path
    
    def _setup_default_styles(self, doc: Document):
        """设置默认样式"""
        # 设置Normal样式
        normal_style = doc.styles['Normal']
        normal_style.font.name = self.default_font
        normal_style.font.size = self.default_font_size
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
        
        # 设置段落格式
        normal_style.paragraph_format.line_spacing = self.line_spacing
        normal_style.paragraph_format.space_after = Pt(6)
    
    def _render_heading(self, doc: Document, item: Dict):
        """渲染标题"""
        title = item.get('title', '')
        level = item.get('level', 1)
        
        # 添加标题
        heading = doc.add_heading(title, level=level)
        
        # 设置标题样式
        heading.style.font.name = self.heading_font
        heading.style.font.size = self.heading_sizes.get(level, Pt(12))
        heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)
        
        # 设置标题格式
        heading.paragraph_format.line_spacing = self.line_spacing
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    def _render_paragraph(self, doc: Document, item: Dict):
        """渲染段落"""
        text = item.get('text', '')
        
        if not text:
            return
        
        # 添加段落
        para = doc.add_paragraph(text)
        
        # 设置段落格式：首行缩进2字符，行距1.5倍
        para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
        para.paragraph_format.line_spacing = self.line_spacing
        para.paragraph_format.space_after = Pt(6)
        
        # 设置字体
        for run in para.runs:
            run.font.name = self.default_font
            run.font.size = self.default_font_size
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
    
    def _render_table(self, doc: Document, item: Dict, tables_data: List[Dict]):
        """渲染表格"""
        table_id = item.get('id')
        
        # 查找表格数据
        table_data = next((t for t in tables_data if t['id'] == table_id), None)
        if not table_data:
            return
        
        content = table_data.get('content')
        if not content:
            return
        
        # 创建表格
        if isinstance(content, list) and len(content) > 0:
            rows = len(content)
            cols = len(content[0]) if content[0] else 1
            
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # 填充表格内容
            for i, row_data in enumerate(content):
                row = table.rows[i]
                for j, cell_data in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = str(cell_data) if cell_data else ''
                    
                    # 设置单元格字体
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = self.line_spacing
                        for run in paragraph.runs:
                            run.font.name = self.default_font
                            run.font.size = Pt(10)  # 表格字体稍小
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
            
            # 表格居中
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        doc.add_paragraph()
    
    def _render_image(self, doc: Document, item: Dict, images_data: List[Dict]):
        """渲染图片"""
        image_id = item.get('id')
        
        # 查找图片数据
        image_data = next((i for i in images_data if i['id'] == image_id), None)
        if not image_data:
            return
        
        content = image_data.get('content')
        if not content:
            return
        
        # 创建临时文件保存图片
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            tmp_file.write(content)
            tmp_path = tmp_file.name
        
        try:
            # 添加图片（居中）
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(tmp_path, width=Inches(4))
            
            # 添加空行
            doc.add_paragraph()
        finally:
            # 删除临时文件
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)