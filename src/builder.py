"""
文档重建模块 - 优化版本
负责根据解析结果和 AI 处理结果重建 Word 文档
优化：按原始 run 顺序重建，保留公式 run 的原始结构
"""
import os
import tempfile
from copy import deepcopy
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.utils import is_document_title_candidate


class DocumentBuilder:
    """文档重建器"""
    
    def __init__(self, elements: List, polished_paragraphs: Dict, new_order: List[str], generated_chart_titles: Dict = None, generated_chapter_titles: Dict = None, document_title: str = None, title_from_first_paragraph: bool = False, chart_title_paragraphs: List = None):
        """
        初始化重建器
        
        Args:
            elements: 元素列表
            polished_paragraphs: 润色后的段落字典
            new_order: 新元素顺序
            generated_chart_titles: AI生成的图表标题 {element_id: title}
            generated_chapter_titles: AI生成的章节标题 {paragraph_id: title}
            document_title: 文档主标题
            title_from_first_paragraph: 首段是否是文档标题
            chart_title_paragraphs: 被识别为图表标题的段落ID列表
        """
        self.elements = elements
        self.polished_paragraphs = polished_paragraphs
        self.new_order = new_order
        self.elem_dict = {elem[1]: elem for elem in elements}
        self.generated_chart_titles = generated_chart_titles or {}
        self.generated_chapter_titles = generated_chapter_titles or {}
        self.document_title = document_title
        self.title_from_first_paragraph = title_from_first_paragraph
        self.chart_title_paragraphs = set(chart_title_paragraphs or [])
    
    def build(self, output_path: str):
        """
        重建文档
        
        Args:
            output_path: 输出文件路径
        """
        doc = Document()
        
        # 添加文档主标题（如果有）
        if self.document_title:
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run(self.document_title)
            title_run.bold = True
            title_run.font.size = Pt(22)
            title_run.font.name = '黑体'
        
        # 如果首段是文档标题，需要跳过
        first_paragraph_id = None
        if self.title_from_first_paragraph:
            # 找到第一个段落ID
            for elem_id in self.new_order:
                if elem_id.startswith('P'):
                    first_paragraph_id = elem_id
                    break
        
        for i, elem_id in enumerate(self.new_order):
            if elem_id not in self.elem_dict:
                continue
            
            # 检查是否有章节标题需要插入（在当前段落之前）
            if elem_id in self.generated_chapter_titles:
                chapter_title = self.generated_chapter_titles[elem_id]
                self._build_chapter_title_paragraph(doc, chapter_title)
            
            elem_type, _, content = self.elem_dict[elem_id]
            
            if elem_type == 'paragraph':
                # 如果首段是文档标题，跳过首段
                if elem_id == first_paragraph_id:
                    continue
                self._build_paragraph(doc, elem_id, content, False)
            
            elif elem_type == 'table':
                # 使用AI生成的表格标题（如果有）
                if elem_id in self.generated_chart_titles:
                    ai_title = self.generated_chart_titles[elem_id]
                    self._build_chart_title_paragraph(doc, ai_title)
                
                self._build_table(doc, content)
            
            elif elem_type == 'image':
                # 先输出图片
                self._build_image(doc, content)
                
                # 使用AI生成的图片标题（放在图片底部）
                if elem_id in self.generated_chart_titles:
                    ai_title = self.generated_chart_titles[elem_id]
                    self._build_chart_title_paragraph(doc, ai_title)
            
            elif elem_type == 'formula':
                self._build_formula(doc, content)
        
        doc.save(output_path)
    
    def _build_paragraph(self, doc: Document, elem_id: str, content: Any, is_chart_title_para: bool = False):
        """重建段落
        
        Args:
            doc: 文档对象
            elem_id: 元素ID
            content: 元素内容
            is_chart_title_para: 是否是图表标题段落
        """
        # 获取文本内容（用于空段落判断）
        text_to_check = ""
        if isinstance(content, tuple) and len(content) >= 1:
            text_to_check = content[0] if content[0] else ""
        elif isinstance(content, str):
            text_to_check = content
        
        # 处理润色后的文本
        if elem_id in self.polished_paragraphs:
            text_to_check = self.polished_paragraphs[elem_id] or text_to_check
        
        # 如果没有文本内容，也不是图表标题段落，则跳过（避免空行）
        if not is_chart_title_para and not text_to_check.strip():
            return
        
        # 处理新的段落格式（元组格式）
        if isinstance(content, tuple) and len(content) >= 5:
            # 解包段落数据
            original_text = content[0]
            formula_runs_list = content[1]
            original_para = content[2]
            all_runs = content[3]
            omath_elements = content[4]
            # 获取包含公式的 run 索引列表（第 6 个元素）
            formula_run_indices = content[5] if len(content) >= 6 else None
            
            # 使用润色后的文字
            if elem_id in self.polished_paragraphs:
                polished_text = self.polished_paragraphs[elem_id]
            else:
                polished_text = original_text
            
            # 创建新段落
            p = doc.add_paragraph()
            
            # 设置段落格式：首行缩进2字符，行距1.5倍
            if not is_chart_title_para:
                # 首行缩进2字符（约0.5cm）
                p.paragraph_format.first_line_indent = Pt(24)  # 24磅约等于2个中文字符
                # 行距1.5倍
                p.paragraph_format.line_spacing = 1.5
            else:
                # 图表标题居中，不缩进
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 如果段落包含公式（包括行内公式和独立公式），需要保留原始 runs
            # 判断是否有公式：检查 original_para 中是否包含 oMath 元素
            has_formula = False
            try:
                omath_check = original_para._element.xpath('.//m:oMath') + original_para._element.xpath('.//*[local-name()="oMath"]')
                has_formula = len(omath_check) > 0
            except:
                has_formula = False
            
            if has_formula and all_runs:
                # 检查是否有行内公式需要处理
                if formula_run_indices:
                    # 处理行内公式：将占位符替换为实际公式
                    omath_list = omath_elements if omath_elements else []
                    self._build_inline_formula_paragraph(p, polished_text, original_para, formula_run_indices, omath_list)
                else:
                    # 原有策略：添加润色后的文本，然后在末尾附加 oMath 元素
                    self._build_formula_at_end_paragraph(p, polished_text, original_para, omath_elements)
            else:
                # 没有公式，直接添加润色后的文本
                run = p.add_run(polished_text)
                
                # 复制原始段落的格式
                if original_para and original_para.runs:
                    first_run = original_para.runs[0]
                    run.bold = first_run.bold
                    run.italic = first_run.italic
                    run.underline = first_run.underline
                    if first_run.font.name:
                        run.font.name = first_run.font.name
                    if first_run.font.size:
                        run.font.size = first_run.font.size
        else:
            # 旧格式处理
            if isinstance(content, tuple):
                original_text = content[0]
            else:
                original_text = content
            if elem_id in self.polished_paragraphs:
                text = self.polished_paragraphs[elem_id]
            else:
                text = original_text
            p = doc.add_paragraph(text)
            
            # 设置段落格式：首行缩进2字符，行距1.5倍
            if not is_chart_title_para:
                # 首行缩进2字符（约0.5cm）
                p.paragraph_format.first_line_indent = Pt(24)
                # 行距1.5倍
                p.paragraph_format.line_spacing = 1.5
            else:
                # 图表标题居中，不缩进
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _build_inline_formula_paragraph(self, p, polished_text: str, original_para, formula_run_indices: List[int], omath_elements: List):
        """
        构建包含行内公式的段落，将占位符替换为实际公式
        
        Args:
            p: 段落对象
            polished_text: 润色后的文本（可能包含 {{FORMULA#1#}} 等占位符）
            original_para: 原始段落对象
            formula_run_indices: 每个 oMath 元素对应的 run 索引列表（-1 表示在段落开头）
            omath_elements: oMath 元素列表
        """
        # 使用正则表达式提取占位符并替换为公式
        import re
        
        # 查找所有占位符 - 支持两种格式：{{FORMULA#数字#}} 和 {FORMULA#数字#}
        placeholder_pattern_double = r'\{\{FORMULA#(\d+)#\}\}'
        placeholder_pattern_single = r'\{FORMULA#(\d+)#\}'
        
        # 先尝试匹配双花括号
        matches = list(re.finditer(placeholder_pattern_double, polished_text))
        
        # 如果没有双花括号占位符，尝试匹配单花括号（容错处理）
        if not matches:
            matches = list(re.finditer(placeholder_pattern_single, polished_text))
        
        if not matches:
            # 没有占位符，直接添加文本
            run = p.add_run(polished_text)
            if original_para.runs:
                self._copy_run_formatting(run, original_para.runs[0])
            return
        
        # 按占位符分割文本并重建
        last_end = 0
        for match in matches:
            # 添加占位符前的文本
            text_before = polished_text[last_end:match.start()]
            if text_before:
                new_run = p.add_run(text_before)
                if original_para.runs:
                    self._copy_run_formatting(new_run, original_para.runs[0])
            
            # 获取公式索引
            formula_idx = int(match.group(1)) - 1  # 转换为 0-based 索引
            
            # 插入对应的 oMath 元素
            if formula_idx < len(omath_elements):
                new_omath = deepcopy(omath_elements[formula_idx])
                p._element.append(new_omath)
            
            last_end = match.end()
        
        # 添加最后的文本
        text_after = polished_text[last_end:]
        if text_after:
            new_run = p.add_run(text_after)
            if original_para.runs:
                self._copy_run_formatting(new_run, original_para.runs[0])
    
    def _split_text_to_segments(self, text: str, num_segments: int) -> List[str]:
        """
        将文本分割成指定数量的段
        
        Args:
            text: 要分割的文本
            num_segments: 段数
            
        Returns:
            文本段列表
        """
        if num_segments <= 0:
            return [text]
        
        if num_segments == 1:
            return [text]
        
        # 简单策略：按字符数平均分割
        total_len = len(text)
        segment_len = total_len // num_segments
        segments = []
        
        for i in range(num_segments):
            start = i * segment_len
            if i == num_segments - 1:
                # 最后一段包含剩余所有文本
                segments.append(text[start:])
            else:
                end = start + segment_len
                segments.append(text[start:end])
        
        return segments
    
    def _build_formula_at_end_paragraph(self, p, polished_text: str, original_para, omath_elements: List):
        """
        构建段落，将公式放在末尾
        
        Args:
            p: 段落对象
            polished_text: 润色后的文本
            original_para: 原始段落对象
            omath_elements: oMath 元素列表
        """
        # 添加润色后的文本
        run = p.add_run(polished_text)
        
        # 复制原始段落的第一个 run 的格式
        if original_para and original_para.runs:
            first_run = original_para.runs[0]
            run.bold = first_run.bold
            run.italic = first_run.italic
            run.underline = first_run.underline
            if first_run.font.name:
                run.font.name = first_run.font.name
            if first_run.font.size:
                run.font.size = first_run.font.size
        
        # 添加所有 oMath 元素到段落末尾，每个公式之间用空格隔开
        if omath_elements:
            for i, omath_elem in enumerate(omath_elements):
                # 在第一个公式前添加空格
                if i == 0:
                    space_run = p.add_run(' ')
                # 添加公式
                new_omath = deepcopy(omath_elem)
                p._element.append(new_omath)
                # 在公式后添加空格（除了最后一个公式）
                if i < len(omath_elements) - 1:
                    space_run = p.add_run(' ')
    
    def _copy_run_formatting(self, run, original_run):
        """复制原始 run 的格式到指定 run"""
        if original_run:
            run.bold = original_run.bold
            run.italic = original_run.italic
            run.underline = original_run.underline
            if original_run.font.name:
                run.font.name = original_run.font.name
            if original_run.font.size:
                run.font.size = original_run.font.size
    
    def _build_table(self, doc: Document, content: List):
        """重建表格"""
        table_data = content
        if table_data and len(table_data) > 0:
            rows = len(table_data)
            cols = len(table_data[0]) if rows > 0 else 1
            try:
                # 创建新表格
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                # 设置表格居中
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 填充表格数据
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < cols and cell is not None:
                            table.cell(i, j).text = str(cell)
            except Exception as e:
                doc.add_paragraph(f"[表格重建失败：{str(e)}]")

    def _build_image(self, doc: Document, content: bytes):
        """重建图片"""
        img_data = content
        if img_data:
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp.write(img_data)
                tmp_path = tmp.name
            try:
                # 创建居中的段落
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 在段落中添加图片
                run = p.add_run()
                run.add_picture(tmp_path, width=Inches(4))
            finally:
                os.unlink(tmp_path)
    
    def _build_formula(self, doc: Document, content: Any):
        """重建公式"""
        if content is not None:
            p = doc.add_paragraph()
            new_formula_run = p.add_run('')
            new_formula_run._element.append(deepcopy(content._element))
        else:
            p = doc.add_paragraph()
            run = p.add_run("[公式]")
            run.italic = True
    
    def _build_chart_title_paragraph(self, doc: Document, title_text: str):
        """构建AI生成的图表标题段落"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        run.font.size = Pt(10.5)
    
    def _build_chapter_title_paragraph(self, doc: Document, title_text: str):
        """构建AI生成的章节标题段落"""
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(15)
