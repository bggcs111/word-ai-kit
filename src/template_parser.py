"""
Word 模板解析器
从 Word 模板文档提取样式、结构和生成模板定义
支持通过 AI 智能推断章节描述
"""
import os
import re
import json
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.template import (
    TemplateDefinition, TemplateElement, StyleDefinition,
    FontDefinition, ParagraphStyle, ElementType
)
from src.logger import log_info, log_error, log_debug
from src.utils import clean_ai_json_response, extract_json_from_text


class WordTemplateParser:
    """Word 模板解析器"""
    
    HEADING_STYLES = {
        'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3,
        'Heading 4': 4, 'Heading 5': 5, 'Heading 6': 6,
        '标题 1': 1, '标题 2': 2, '标题 3': 3,
        '标题 4': 4, '标题 5': 5, '标题 6': 6,
    }
    
    def __init__(self):
        self.template_doc: Optional[Document] = None
        self.styles_map: Dict[str, StyleDefinition] = {}
        self.structure: List[TemplateElement] = []
        self.template_def: Optional[TemplateDefinition] = None
    
    def parse_template(
        self,
        template_path: str,
        ai_client=None,
        model_name: str = ""
    ) -> TemplateDefinition:
        """
        解析 Word 模板文档
        
        Args:
            template_path: 模板文件路径
            ai_client: OpenAI 客户端实例，用于 AI 推断章节描述
            model_name: AI 模型名称
            
        Returns:
            TemplateDefinition 对象
        """
        log_info(f"开始解析模板：{template_path}")
        
        self.template_doc = Document(template_path)
        
        template_name = Path(template_path).stem
        self.template_def = TemplateDefinition(
            template_id=template_name,
            name=template_name,
            description=f"从 {Path(template_path).name} 自动提取的模板"
        )
        
        self._extract_styles()
        self._extract_structure()
        
        if ai_client and model_name:
            heading_contents = self._extract_heading_content()
            self._ai_infer_descriptions(ai_client, model_name, heading_contents)
            log_info("AI 章节描述推断完成")
        else:
            log_info("未提供 AI 客户端，跳过章节描述推断")
        
        log_info(f"模板解析完成：{len(self.styles_map)} 个样式，{len(self.structure)} 个元素")
        
        return self.template_def
    
    def _extract_styles(self):
        """从模板文档提取所有使用的样式"""
        doc = self.template_doc
        
        for para in doc.paragraphs:
            style_name = para.style.name
            
            if style_name not in self.styles_map:
                self.styles_map[style_name] = self._parse_style(para)
    
    def _parse_style(self, para) -> StyleDefinition:
        """解析段落的样式
        
        改进：遍历所有 run 找到有字体设置的 run，同时回退到样式级别读取格式
        """
        style_name = para.style.name
        
        font_def = self._parse_font_from_paragraph(para)
        para_style = self._parse_paragraph_format(para)
        
        is_heading = False
        heading_level = None
        is_title = False
        
        if style_name in self.HEADING_STYLES:
            is_heading = True
            heading_level = self.HEADING_STYLES[style_name]
        
        if style_name.lower() in ['title', '标题']:
            is_title = True
        
        return StyleDefinition(
            style_name=style_name,
            font=font_def,
            paragraph=para_style,
            is_heading=is_heading,
            heading_level=heading_level,
            is_title=is_title
        )
    
    def _parse_font_from_paragraph(self, para) -> FontDefinition:
        """从段落解析字体，遍历所有 run 并回退到样式级别
        
        策略：
        1. 先读取样式级别的字体作为基础
        2. 遍历所有 run，找到有显式字体设置的 run
        3. 合并 run 级别和样式级别的设置
        """
        font = FontDefinition()
        
        style_font_name = None
        style_font_size = None
        style_font_bold = None
        style_font_italic = None
        style_font_underline = None
        style_font_color = None
        
        try:
            style_font = para.style.font
            style_font_name = style_font.name
            style_font_size = style_font.size
            style_font_bold = style_font.bold
            style_font_italic = style_font.italic
            style_font_underline = style_font.underline
            if style_font.color and style_font.color.rgb:
                style_font_color = str(style_font.color.rgb)
        except:
            pass
        
        run_font_name = None
        run_font_size = None
        run_font_bold = None
        run_font_italic = None
        run_font_underline = None
        run_font_color = None
        
        for run in para.runs:
            if run.font.name and not run_font_name:
                run_font_name = run.font.name
            if run.font.size and not run_font_size:
                run_font_size = run.font.size
            if run.bold is not None and run_font_bold is None:
                run_font_bold = run.bold
            if run.italic is not None and run_font_italic is None:
                run_font_italic = run.italic
            if run.underline is not None and run_font_underline is None:
                run_font_underline = run.underline
            if run.font.color and run.font.color.rgb and not run_font_color:
                try:
                    run_font_color = str(run.font.color.rgb)
                except:
                    pass
        
        final_name = run_font_name or style_font_name
        final_size = run_font_size or style_font_size
        final_bold = run_font_bold if run_font_bold is not None else style_font_bold
        final_italic = run_font_italic if run_font_italic is not None else style_font_italic
        final_underline = run_font_underline if run_font_underline is not None else style_font_underline
        final_color = run_font_color or style_font_color
        
        if final_name:
            font.name = final_name
        if final_size:
            font.size = f"{final_size.pt}pt"
        if final_bold is not None:
            font.bold = bool(final_bold)
        if final_italic is not None:
            font.italic = bool(final_italic)
        if final_underline is not None:
            font.underline = bool(final_underline)
        if final_color:
            font.color = final_color
        
        return font
    
    def _parse_paragraph_format(self, para) -> ParagraphStyle:
        """解析段落格式，支持从段落直接格式和样式级别回退读取
        
        策略：
        1. 先读取段落直接格式
        2. 如果为 None，回退到样式级别的段落格式
        """
        pf = ParagraphStyle()
        
        fmt = para.paragraph_format
        
        style_fmt = None
        try:
            style_fmt = para.style.paragraph_format
        except:
            pass
        
        alignment = fmt.alignment
        if alignment is None and style_fmt:
            alignment = style_fmt.alignment
        if alignment:
            align_map = {
                WD_ALIGN_PARAGRAPH.LEFT: 'left',
                WD_ALIGN_PARAGRAPH.CENTER: 'center',
                WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
            }
            pf.alignment = align_map.get(alignment, 'left')
        
        line_spacing = fmt.line_spacing
        if line_spacing is None and style_fmt:
            line_spacing = style_fmt.line_spacing
        if line_spacing:
            pf.line_spacing = float(line_spacing)
        elif fmt.line_spacing_rule is not None or (style_fmt and style_fmt.line_spacing_rule is not None):
            pf.line_spacing = 1.5
        
        space_before = fmt.space_before
        if space_before is None and style_fmt:
            space_before = style_fmt.space_before
        if space_before:
            pf.space_before = f"{space_before.pt}pt"
        
        space_after = fmt.space_after
        if space_after is None and style_fmt:
            space_after = style_fmt.space_after
        if space_after:
            pf.space_after = f"{space_after.pt}pt"
        
        first_line_indent = fmt.first_line_indent
        if first_line_indent is None and style_fmt:
            first_line_indent = style_fmt.first_line_indent
        if first_line_indent:
            pf.first_line_indent = f"{first_line_indent.cm}cm"
        
        left_indent = fmt.left_indent
        if left_indent is None and style_fmt:
            left_indent = style_fmt.left_indent
        if left_indent:
            pf.indent_left = f"{left_indent.cm}cm"
        
        right_indent = fmt.right_indent
        if right_indent is None and style_fmt:
            right_indent = style_fmt.right_indent
        if right_indent:
            pf.indent_right = f"{right_indent.cm}cm"
        
        return pf
    
    def _extract_structure(self):
        """提取文档结构"""
        order = 0
        
        for para in self.template_doc.paragraphs:
            style_name = para.style.name
            text = para.text.strip()
            
            elem = self._create_element(para, style_name, text, order)
            if elem:
                self.structure.append(elem)
                order += 1
        
        for table in self.template_doc.tables:
            elem = TemplateElement(
                element_type=ElementType.TABLE,
                style_name="Table",
                placeholder_text=f"[{table.rows.__len__()}行 x {len(table.columns)}列表格]",
                order=order
            )
            self.structure.append(elem)
            order += 1
        
        self.template_def.structure = self.structure
        self.template_def.styles = self.styles_map
    
    def _create_element(self, para, style_name: str, text: str, order: int) -> Optional[TemplateElement]:
        """创建模板元素"""
        if style_name in self.HEADING_STYLES:
            level = self.HEADING_STYLES[style_name]
            numbering = self._extract_numbering(text)
            clean_text = self._remove_numbering(text)
            
            return TemplateElement(
                element_type=ElementType.HEADING,
                style_name=style_name,
                heading_level=level,
                numbering=numbering,
                placeholder_text=clean_text or f"第{level}级标题",
                description="",
                order=order
            )
        
        if style_name.lower() in ['title', '标题']:
            return TemplateElement(
                element_type=ElementType.TITLE,
                style_name=style_name,
                placeholder_text=text or "文档标题",
                order=order
            )
        
        if text.startswith('[图片]') or text.startswith('[IMAGE]'):
            return TemplateElement(
                element_type=ElementType.IMAGE,
                style_name=style_name,
                placeholder_text=text,
                order=order
            )
        
        if text.startswith('[表格]') or text.startswith('[TABLE]'):
            return TemplateElement(
                element_type=ElementType.TABLE,
                style_name=style_name,
                placeholder_text=text,
                order=order
            )
        
        return TemplateElement(
            element_type=ElementType.PARAGRAPH,
            style_name=style_name,
            placeholder_text=text[:50] if text else "[正文内容]",
            order=order
        )
    
    def _extract_numbering(self, text: str) -> Optional[str]:
        """提取编号"""
        patterns = [
            r'^(\d+\.\d+\.\d+)\s',
            r'^(\d+\.\d+)\s',
            r'^(\d+)\s',
            r'^第[一二三四五六七八九十]+[章节部分]\s',
            r'^([A-Z]\.)\s',
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                return match.group(1)
        
        return None
    
    def _remove_numbering(self, text: str) -> str:
        """移除编号"""
        patterns = [
            r'^(\d+\.\d+\.\d+)\s*',
            r'^(\d+\.\d+)\s*',
            r'^(\d+)\s*',
            r'^第[一二三四五六七八九十]+[章节部分]\s*',
            r'^([A-Z]\.)\s*',
        ]
        
        for pattern in patterns:
            text = re.sub(pattern, '', text)
        
        return text.strip()
    
    def _extract_heading_content(self) -> Dict[str, str]:
        """
        提取每个标题下的段落内容
        
        遍历文档段落，将每个标题到下一个标题之间的段落文本
        收集为该标题的内容摘要
        
        Returns:
            字典：{标题文本: 内容摘要}
        """
        heading_contents: Dict[str, str] = {}
        current_heading = None
        current_texts: List[str] = []
        max_content_length = 300
        
        for para in self.template_doc.paragraphs:
            style_name = para.style.name
            text = para.text.strip()
            
            if not text:
                continue
            
            if style_name in self.HEADING_STYLES or style_name.lower() in ['title', '标题']:
                if current_heading and current_texts:
                    combined = ' '.join(current_texts)
                    if len(combined) > max_content_length:
                        combined = combined[:max_content_length] + "..."
                    heading_contents[current_heading] = combined
                
                clean_text = self._remove_numbering(text)
                current_heading = clean_text or text
                current_texts = []
            else:
                if current_heading:
                    current_texts.append(text)
        
        if current_heading and current_texts:
            combined = ' '.join(current_texts)
            if len(combined) > max_content_length:
                combined = combined[:max_content_length] + "..."
            heading_contents[current_heading] = combined
        
        return heading_contents
    
    def _ai_infer_descriptions(
        self,
        ai_client,
        model_name: str,
        heading_contents: Dict[str, str]
    ):
        """
        通过 AI 批量推断章节描述
        
        将所有标题及其内容摘要一次性发送给 AI，
        让 AI 理解每个章节的用途并生成描述
        
        Args:
            ai_client: OpenAI 客户端
            model_name: 模型名称
            heading_contents: {标题文本: 内容摘要}
        """
        if not heading_contents:
            log_info("没有标题内容可供 AI 推断")
            return
        
        sections_text = ""
        for i, (heading, content) in enumerate(heading_contents.items(), 1):
            sections_text += f"\n{i}. 章节名称：{heading}"
            if content:
                sections_text += f"\n   内容摘要：{content}"
            else:
                sections_text += f"\n   内容摘要：（无内容，为空章节）"
        
        prompt = f"""你是一个专业的文档结构分析专家。请根据以下模板文档的章节名称和每个章节的内容摘要，推断每个章节的用途描述。

要求：
1. 描述应简洁明了，说明该章节应该包含什么类型的内容
2. 描述应从文档使用者的角度出发，指导后续内容生成
3. 如果章节下没有内容（空章节），根据章节名称合理推断其用途
4. 只输出 JSON，不要有任何额外说明

## 模板章节结构
{sections_text}

## 输出格式
请严格按照以下 JSON 格式输出，key 为章节名称，value 为推断的用途描述：
{{"章节名称1": "用途描述1", "章节名称2": "用途描述2"}}"""
        
        try:
            response = ai_client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": "你是专业的文档结构分析专家，擅长理解文档章节的用途和内容定位。只输出 JSON 格式。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
            )
            
            result = response.choices[0].message.content.strip()
            descriptions = self._parse_ai_descriptions(result)
            
            if descriptions:
                self._apply_descriptions(descriptions)
                log_info(f"AI 推断了 {len(descriptions)} 个章节描述")
            else:
                log_error("AI 描述推断结果解析失败")
            
        except Exception as e:
            log_error(f"AI 章节描述推断失败：{str(e)}")
    
    def _parse_ai_descriptions(self, result: str) -> Dict[str, str]:
        """
        解析 AI 返回的描述 JSON（使用工具函数）
        
        Args:
            result: AI 原始返回文本
            
        Returns:
            {标题文本: 描述} 字典
        """
        data = extract_json_from_text(result)
        if not data:
            log_error(f"无法从 AI 响应中提取 JSON: {result[:200]}")
            return {}
        
        try:
            if isinstance(data, dict):
                return {str(k): str(v) for k, v in data.items()}
            return {}
        except Exception as e:
            log_error(f"AI 描述 JSON 解析失败：{e}")
            return {}
    
    def _apply_descriptions(self, descriptions: Dict[str, str]):
        """
        将 AI 推断的描述应用到模板结构中
        
        通过 placeholder_text 匹配标题元素，填入对应描述
        
        Args:
            descriptions: {标题文本: 描述} 字典
        """
        for elem in self.structure:
            if elem.element_type == ElementType.HEADING:
                desc = descriptions.get(elem.placeholder_text, "")
                if desc:
                    elem.description = desc
    
    def save_template_json(self, output_path: str) -> str:
        """保存模板为 JSON 文件"""
        if not self.template_def:
            raise ValueError("未解析模板")
        
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.template_def.to_dict(), f, ensure_ascii=False, indent=2)
        
        log_info(f"模板 JSON 已保存：{output_path}")
        return output_path
    
    @staticmethod
    def load_template_json(json_path: str) -> TemplateDefinition:
        """从 JSON 文件加载模板"""
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return TemplateDefinition.from_dict(data)
    
    @staticmethod
    def get_templates_dir() -> Path:
        """获取模板目录"""
        templates_dir = Path(__file__).parent.parent / "templates"
        templates_dir.mkdir(exist_ok=True)
        return templates_dir
    
    @staticmethod
    def list_saved_templates() -> List[Dict]:
        """列出所有已保存的模板"""
        templates_dir = WordTemplateParser.get_templates_dir()
        templates = []
        
        for json_file in templates_dir.glob("*.json"):
            try:
                with open(json_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                templates.append({
                    "template_id": data.get("template_id", json_file.stem),
                    "name": data.get("name", json_file.stem),
                    "description": data.get("description", ""),
                    "version": data.get("version", "1.0"),
                    "file_path": str(json_file),
                    "structure_count": len(data.get("structure", [])),
                    "styles_count": len(data.get("styles", {}))
                })
            except Exception as e:
                log_error(f"加载模板失败 {json_file}: {e}")
        
        return templates
