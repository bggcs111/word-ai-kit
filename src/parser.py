"""
文档解析模块 - 优化版本
负责解析 Word 文档，提取段落、表格、图片、公式等元素
优化：记录每个 run 是否包含公式，重建时保留公式 run 的原始结构
"""
import re
from typing import List, Dict, Tuple, Any
from docx import Document
from src.constants import INDEPENDENT_FORMULA_TEXT_THRESHOLD


class DocumentParser:
    """文档解析器"""
    
    def __init__(self):
        """初始化解析器"""
        self.elements = []  # [(type, id, content), ...]
        # 段落数据：(text, formula_runs, original_para, all_runs, omath_elements, formula_run_indices)
        # formula_run_indices: 包含公式的 run 索引列表
        self.paragraphs = {}
        self.p_count = 1
        self.t_count = 1
        self.i_count = 1
        self.f_count = 1
    
    def parse(self, doc_path: str) -> Tuple[List, Dict, Document]:
        """
        解析 Word 文档
        
        Args:
            doc_path: 文档路径
            
        Returns:
            (elements, paragraphs, original_doc)
        """
        doc = Document(doc_path)
        self.elements = []
        self.paragraphs = {}
        self.p_count = 1
        self.t_count = 1
        self.i_count = 1
        self.f_count = 1
        
        # 获取文档的 body 元素，按顺序遍历段落和表格
        body = doc.element.body
        
        # 遍历 body 中的所有子元素（包括段落和表格）
        for child in body.iterchildren():
            if child.tag.endswith('p'):  # 段落
                self._parse_paragraph(child, doc)
            elif child.tag.endswith('tbl'):  # 表格
                self._parse_table(child, doc)
        
        return self.elements, self.paragraphs, doc
    
    def _parse_paragraph(self, para_element, doc: Document):
        """解析段落"""
        # 找到对应的 paragraph 对象
        para = None
        for p in doc.paragraphs:
            if p._element == para_element:
                para = p
                break
        
        if para is None:
            return
        
        text = para.text.strip() if para.text else ""
        
        # 检查图片
        img_data_list = self._extract_images(para, doc)
        
        # 检查公式
        has_formula, has_inline_formula, formula_runs_list, all_runs_with_formulas, omath_elements = \
            self._extract_formulas(para, text)
        
        # 添加段落文字（包含行内公式的段落）
        # 注意：独立公式段落不添加为段落元素，只作为公式元素添加
        if (text or img_data_list) or has_inline_formula:
            pid = f"P{self.p_count}"
            
            # 对于行内公式，记录哪些 run 包含公式
            formula_run_indices = None
            if has_inline_formula:
                formula_run_indices = self._get_formula_run_indices(para)
            
            # 保存段落信息
            para_data = (text, formula_runs_list if has_inline_formula else None, para, all_runs_with_formulas, omath_elements if has_inline_formula else None, formula_run_indices)
            self.paragraphs[pid] = para_data
            self.elements.append(('paragraph', pid, para_data))
            self.p_count += 1
            
            # 调试信息
            if has_formula:
                print(f"段落 {pid} 检测到公式：{len(all_runs_with_formulas)} 个 runs, {len(formula_runs_list)} 个公式 runs, {len(omath_elements)} 个 oMath 元素，独立公式={not has_inline_formula}")
                if formula_run_indices:
                    print(f"  公式 run 索引：{formula_run_indices}")
        
        # 添加图片（作为独立元素，紧跟在相关段落后）
        for img_data in img_data_list:
            if img_data:
                iid = f"I{self.i_count}"
                self.elements.append(('image', iid, img_data))
                self.i_count += 1
        
        # 添加独立公式（非行内公式）
        if has_formula and not has_inline_formula:
            for formula_run in formula_runs_list:
                fid = f"F{self.f_count}"
                self.elements.append(('formula', fid, formula_run))
                self.f_count += 1
                print(f"添加独立公式：{fid}")
    
    def _parse_table(self, table_element, doc: Document):
        """解析表格"""
        # 找到对应的 table 对象
        table = None
        for tbl in doc.tables:
            if tbl._element == table_element:
                table = tbl
                break
        
        if table is None:
            return
        
        tid = f"T{self.t_count}"
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        self.elements.append(('table', tid, table_data))
        print(f"添加表格：{tid} (在当前位置)")
        self.t_count += 1
    
    def _extract_images(self, para, doc: Document) -> List[bytes]:
        """提取段落中的图片"""
        img_data_list = []
        for run in para.runs:
            blips = run._element.xpath('.//a:blip')
            for blip in blips:
                rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rid:
                    image_part = doc.part.related_parts[rid]
                    img_data_list.append(image_part.blob)
        return img_data_list
    
    def _extract_formulas(self, para, text: str) -> Tuple[bool, bool, List, List, List]:
        """
        提取段落中的公式
        
        Returns:
            (has_formula, has_inline_formula, formula_runs_list, all_runs_with_formulas, omath_elements)
        """
        formula_runs_list = []
        all_runs_with_formulas = []
        has_formula = False
        has_inline_formula = False
        omath_elements_list = []  # 保存 oMath 元素本身
        
        # 使用更安全的公式检测方法
        try:
            # 尝试标准 Office Math 命名空间
            omath_elements = para._element.xpath('.//m:oMath')
        except:
            omath_elements = []
        
        try:
            # 尝试直接查找 oMath 元素（不指定命名空间）
            any_omath_elements = para._element.xpath('.//*[local-name()="oMath"]')
        except:
            any_omath_elements = []
        
        # 检查段落中是否包含公式对象
        if omath_elements or any_omath_elements:
            has_formula = True
            all_omath = omath_elements + any_omath_elements
            # 保存 oMath 元素（去重）
            seen = set()
            for elem in all_omath:
                elem_id = id(elem)
                if elem_id not in seen:
                    omath_elements_list.append(elem)
                    seen.add(elem_id)
            
            # 判断是行内公式还是独立公式
            # 独立公式的特征：段落只有公式，没有文字或文字极少
            text_len = len(text) if text else 0
            
            # 如果段落只有公式（没有文字或文字极少），则是独立公式
            if text_len < INDEPENDENT_FORMULA_TEXT_THRESHOLD and len(all_omath) > 0:
                # 独立公式 - 直接保存整个段落的所有 runs，确保公式不丢失
                has_inline_formula = False
                all_runs_with_formulas = list(para.runs)
                # 独立公式直接使用所有 runs 作为公式 runs
                # 如果没有 runs（公式直接在段落级别），保存段落对象
                if para.runs:
                    formula_runs_list = list(para.runs)
                else:
                    # 公式在段落级别，保存段落对象
                    formula_runs_list = [para]
                print(f"  -> 独立公式：{len(all_omath)} 个 oMath 元素，{len(formula_runs_list)} 个 runs/para")
            else:
                # 行内公式 - 文字和公式混合
                has_inline_formula = True
                # 对于行内公式，保存整个段落的 runs 和 oMath 元素
                all_runs_with_formulas = list(para.runs)
                formula_runs_list = list(para.runs)
        
        return has_formula, has_inline_formula, formula_runs_list, all_runs_with_formulas, omath_elements_list
    
    def _get_formula_run_indices(self, para) -> List[int]:
        """
        获取每个 oMath 元素应该在哪个 run 之后插入
        通过分析 oMath 元素在段落 XML 中的位置来确定
        
        Args:
            para: 段落对象
            
        Returns:
            每个 oMath 元素对应的 run 索引列表（索引从 0 开始，-1 表示在段落开头）
        """
        formula_run_indices = []
        
        try:
            # 获取段落的所有子元素（按顺序）
            children = list(para._element.iterchildren())
            
            # 获取所有 oMath 元素
            omath_elements = para._element.xpath('.//m:oMath') + para._element.xpath('.//*[local-name()="oMath"]')
            
            # 去重
            seen = set()
            unique_omath = []
            for elem in omath_elements:
                elem_id = id(elem)
                if elem_id not in seen:
                    unique_omath.append(elem)
                    seen.add(elem_id)
            omath_elements = unique_omath
            
            print(f"  段落 XML 结构：共 {len(children)} 个子元素")
            
            # 对于每个 oMath 元素，找到它前面的 run 数量
            for omath_idx, omath in enumerate(omath_elements):
                run_count_before = 0
                for child in children:
                    # 如果当前元素是 oMath，停止计数
                    if child == omath:
                        break
                    # 统计 oMath 之前的 run 数量
                    if child.tag.endswith('r'):  # run 元素
                        run_count_before += 1
                
                # run_count_before 表示 oMath 前面有多少个 run
                # 如果 run_count_before >= run 总数，说明 oMath 在段落末尾
                # 我们需要将其转换为 run 索引：run_count_before - 1
                total_runs = len(para.runs)
                if run_count_before >= total_runs:
                    # oMath 在最后一个 run 之后
                    formula_run_indices.append(total_runs - 1)
                    print(f"    oMath {omath_idx + 1} 位置：在最后一个 run（索引 {total_runs - 1}）之后")
                elif run_count_before == 0:
                    # oMath 在第一个 run 之前
                    formula_run_indices.append(-1)
                    print(f"    oMath {omath_idx + 1} 位置：在第一个 run 之前")
                else:
                    # oMath 在 run_count_before-1 索引的 run 之后
                    formula_run_indices.append(run_count_before - 1)
                    print(f"    oMath {omath_idx + 1} 位置：在 run {run_count_before - 1} 之后（共 {total_runs} 个 runs）")
            
        except Exception as e:
            print(f"  获取公式 run 索引失败：{e}")
            import traceback
            traceback.print_exc()
        
        return formula_run_indices
